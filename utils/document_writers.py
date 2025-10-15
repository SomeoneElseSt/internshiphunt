"""Document creation utilities for generating downloadable application materials."""

from io import BytesIO
from typing import Dict
import docx

def create_docx_cover_letter(cover_letter: str, profile: Dict = None) -> bytes:
    """Create a DOCX version of the cover letter with optional profile header."""
    doc = docx.Document()

    if profile and 'resume_data' in profile:
        resume_data = profile['resume_data']
        if 'name' in resume_data:
            doc.add_paragraph(resume_data['name'])
        if 'contact' in resume_data:
            doc.add_paragraph(resume_data['contact'])
        if 'address' in resume_data:
            doc.add_paragraph(resume_data['address'])
        doc.add_paragraph()

    paragraphs = cover_letter.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)

    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()
