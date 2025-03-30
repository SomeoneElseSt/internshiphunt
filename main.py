import streamlit as st
import google.generativeai as genai
import base64
import asyncio
import requests
from agno.agent import Agent
from agno.models.google import Gemini
from agno.tools import Toolkit
import PyPDF2 
import docx
import json
import time
import os
import agentql
from typing import Dict, List, Any
from dotenv import load_dotenv
import aiohttp

def create_docx_cover_letter(cover_letter: str, company: str, role: str, profile: Dict = None) -> bytes:
    """Create a DOCX version of the cover letter"""
    doc = docx.Document()

    # Add header with profile info if available
    if profile and 'resume_data' in profile:
        resume_data = profile['resume_data']
        if 'name' in resume_data:
            doc.add_paragraph(resume_data['name'])
        if 'contact' in resume_data:
            doc.add_paragraph(resume_data['contact'])
        if 'address' in resume_data:
            doc.add_paragraph(resume_data['address'])
        doc.add_paragraph()  # Add space

    # Add content
    paragraphs = cover_letter.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)

    # Save to bytes
    from io import BytesIO
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()

def generate_cover_letter(profile: Dict, job: Dict) -> str:
    """Generate a cover letter using Gemini based on profile and job details"""
    prompt = f"""Generate a professional cover letter for this job:
    Company: {job.get('company', 'N/A')}
    Role: {job.get('role', 'N/A')}
    Location: {job.get('location', 'N/A')}

    Using this candidate's profile:
    Skills: {profile.get('resume_data', {}).get('skills', [])}
    Experiences: {profile.get('resume_data', {}).get('experiences', [])}
    Interests: {profile.get('resume_data', {}).get('interests', [])}

    Write a concise, professional cover letter highlighting relevant skills and experiences.
    Keep it under 400 words and follow standard cover letter format."""

    response = model.generate_content(prompt)
    return response.text

def analyze_application_link(url: str) -> tuple[Dict[str, bool], bool]:
    """
    Send application link to AgentQL for analysis of document requirements.
    Returns dict with cover_letter and research_statement requirements.
    """
    print(f"\n--- AgentQL Analysis for URL: {url} ---")
    try:
        api_key = os.environ.get("AGENT_QL_API_KEY")
        if not api_key:
            print("Warning: AGENT_QL_API_KEY not found in environment")
            return {"written_requirements": []}, True

        headers = {
            "X-API-Key": api_key,
            "Content-Type": "application/json"
        }

        payload = {
            "url": url,
            "query": """
                {
                  written_requirements {
                    accepts_cover_letter(True/False)
                    accepts_research_statement(True/False)
                    accepts_why_us_statement(True/False)
                  }
                }
            """,
            "params": {
                "mode": "standard",
                "wait_for": 15,
                "is_scroll_to_bottom_enabled": True
            }
        }

        response = requests.post(
            "https://api.agentql.com/v1/query-data",
            json=payload,
            headers=headers,
            timeout=25
        )

        print(f"===AgentQL Response Status: {response.status_code}===")

        if response.status_code == 200:
            response_json = response.json()
            print("\n=== AgentQL Response Analysis ===")
            print(f"Full Response: {response_json}")

            # Extract the requirement values from the response
            written_reqs = response_json.get("data", {}).get("written_requirements", {})
            
            # Helper function to convert string values to boolean
            def convert_to_bool(value):
                if value == 'True':
                    return True
                elif value == 'False':
                    return False
                elif value == 'null' or value is None:
                    return False
                else:
                    # Try to interpret as a boolean if it's not a recognized string
                    return bool(value)
            
            # Convert string values to proper booleans
            requires_cover_letter = convert_to_bool(written_reqs.get("accepts_cover_letter"))
            requires_research_statement = convert_to_bool(written_reqs.get("accepts_research_statement"))
            requires_why_us_statement = convert_to_bool(written_reqs.get("accepts_why_us_statement"))

            print(f"\nCover Letter Required: {requires_cover_letter}")
            print(f"Research Statement Required: {requires_research_statement}")
            print(f"Why Us Statement Required: {requires_why_us_statement}")

            result = {
                "requires_cover_letter": requires_cover_letter,
                "requires_research_statement": requires_research_statement,
                "requires_why_us_statement": requires_why_us_statement
            }

            print("=" * 30 + "\n")
            return result, True
        else:
            if 'speedyapply' in url and response.status_code == 404:
                return None, False
            st.warning(f"AgentQL request failed for {url}: {response.status_code}")
            return {"written_requirements": []}, True

    except Exception as e:
        st.error(f"Error analyzing application link {url}: {str(e)}")
        return {"written_requirements": []}, True

def generate_research_statement(profile: Dict, job: Dict) -> str:
    """Generate a research statement using Gemini based on profile and job details"""
    prompt = f"""Generate a research statement for this job application:
    Company: {job.get('company', 'N/A')}
    Role: {job.get('role', 'N/A')}
    Location: {job.get('location', 'N/A')}

    Using this candidate's profile:
    Skills: {profile.get('resume_data', {}).get('skills', [])}
    Experiences: {profile.get('resume_data', {}).get('experiences', [])}
    Interests: {profile.get('resume_data', {}).get('interests', [])}

    Write a concise, professional research statement highlighting relevant research experiences, skills, and interests that align with the job role and company.
    Focus on past research projects, methodologies, and outcomes, and how they make the candidate suitable for a research-oriented role.
    Keep it under 500 words and follow a standard research statement format."""

    response = model.generate_content(prompt)
    return response.text

def generate_why_us_statement(profile: Dict, job: Dict) -> str:
    """Generate a 'Why Us' statement using Gemini based on profile and job details"""
    prompt = f"""Generate a 'Why Us' statement for this job application:
    Company: {job.get('company', 'N/A')}
    Role: {job.get('role', 'N/A')}
    Location: {job.get('location', 'N/A')}

    Using this candidate's profile:
    Skills: {profile.get('resume_data', {}).get('skills', [])}
    Experiences: {profile.get('resume_data', {}).get('experiences', [])}
    Interests: {profile.get('resume_data', {}).get('interests', [])}
    Career Goals: {profile.get('responses', {}).get('q_4', 'N/A')} # Assuming Q4 is about career goals

    Write a compelling 'Why Us' statement explaining why the candidate is interested in working for this specific company and in this role.
    Highlight alignment between the candidate's career goals, skills, and the company's mission, values, and opportunities.
    Mention specific aspects of the company or role that are particularly appealing to the candidate, based on their profile and stated career interests.
    Keep it under 300 words and maintain a professional yet enthusiastic tone."""

    response = model.generate_content(prompt)
    return response.text

class GithubTool(Toolkit):
    def __init__(self):
        super().__init__()
        self.name = "github"
        self.headers = {}
        self.session = None
        self.repo_links = [
            {
                "owner": "speedyapply",
                "repo": "2025-AI-College-Jobs",
                "branch": "main",
                "path": "README.md"
            },
            {
                "owner": "SimplifyJobs",
                "repo": "Summer2025-Internships",
                "branch": "main",
                "path": "README.md"
            }
        ]
        self.params = {
            "get_repository": True,
            "search_repositories": True,
            "list_repositories": True
        }
        self.agent = Agent(
            model=Gemini(
                id="gemini-2.0-pro",
                api_key=GOOGLE_API_KEY
            ),
            description="You are an expert at analyzing internship opportunities and matching them to candidate profiles.",
            markdown=True
        )
        try:
            github_key = os.getenv("GITHUB_KEY")
            if github_key:
                self.headers = {"Authorization": f"token {github_key}"}
        except:
            st.warning("No Github key found. Using unauthenticated access (rate limits may apply).")

    async def fetch_readme(self, repo_owner: str, repo_name: str, branch: str = "main", path: str = "README.md") -> Dict[str, Any]:
        """Fetches a README file from a Github repository using get_repository"""
        if self.session is None:
            self.session = aiohttp.ClientSession()

        try:
            # First get repository details to verify access and existence
            repo_url = f"https://api.github.com/repos/{repo_owner}/{repo_name}"
            async with self.session.get(repo_url, headers=self.headers) as repo_response:
                if repo_response.status != 200:
                    print(f"Repository not found or inaccessible: {repo_url}")
                    return {"content": None, "status": "error", "message": "Repository not found"}

                # Now fetch the README content
                content_url = f"https://api.github.com/repos/{repo_owner}/{repo_name}/contents/{path}?ref={branch}"
                async with self.session.get(content_url, headers=self.headers) as response:
                    if response.status == 200:
                        data = await response.json()
                        if 'content' in data:
                            import base64
                            content = base64.b64decode(data['content']).decode('utf-8')
                            print(f"Successfully fetched content from {content_url}")
                            return {"content": content, "status": "success"}
                    print(f"Failed to fetch content: {response.status}")
                    return {"content": None, "status": "error", "message": f"Failed to fetch content: {response.status}"}
        except Exception as e:
            print(f"Error fetching from Github: {e}")
            return {"content": None, "status": "error", "message": str(e)}

    async def close_session(self):
        """Close the aiohttp session"""
        if self.session:
            await self.session.close()
            self.session = None

    def parse_internship_data(self, content: str) -> List[Dict[str, str]]:
        """Parse internship listings from markdown content"""
        internships = []
        lines = content.split('\n')
        current_internship = {}

        for line in lines:
            if '|' in line:
                parts = [p.strip() for p in line.split('|')]
                if len(parts) >= 4:
                    if all(p.startswith('---') for p in parts if p):  # Skip table headers
                        continue
                    current_internship = {
                        "company": parts[1] if len(parts) > 1 else "",
                        "role": parts[2] if len(parts) > 2 else "",
                        "location": parts[3] if len(parts) > 3 else "",
                        "application": parts[4] if len(parts) > 4 else ""
                    }
                    if current_internship["company"] and current_internship["role"]:
                        internships.append(current_internship)

        return internships

    async def fetch_internship_data(self) -> Dict[str, Any]:
        """Fetches and processes internship data from multiple sources in parallel"""
        
        async def fetch_and_parse_repo(repo):
            """Helper function to fetch and parse a single repository"""
            data = await self.fetch_readme(
                repo["owner"],
                repo["repo"],
                repo["branch"],
                repo["path"]
            )
            if data["status"] == "success":
                internships = self.parse_internship_data(data["content"])
                return [{**i, "source": repo["owner"]} for i in internships]
            return []
        
        # Use asyncio.gather to fetch all repositories in parallel
        repo_results = await asyncio.gather(
            *[fetch_and_parse_repo(repo) for repo in self.repo_links],
            return_exceptions=True
        )
        
        # Process results, handling any exceptions
        all_internships = []
        for result in repo_results:
            if isinstance(result, Exception):
                print(f"Error fetching repository: {result}")
                continue
            all_internships.extend(result)
        
        print("Fetched internships from GitHub:", all_internships)
        return {
            "status": "success", 
            "internships": all_internships,
            "total_count": len(all_internships)
        }

    def recommend_internships(self, internships: List[Dict], profile: Dict) -> List[Dict]:
        """Use custom Gemini agent to recommend internships based on profile"""
        recommendation_agent = Agent(
            model=Gemini(
                id='gemini-1.5-pro-002', 
                api_key=GOOGLE_API_KEY
            ),
            description="You are an expert at matching internship opportunities to candidate profiles, with deep understanding of tech industry requirements and career progression.",
            markdown=True
        )

        prompt = f"""Given this candidate profile:
        {json.dumps(profile)}

        And these internships:
        {json.dumps(internships)}

        Go through the available internships and return at maximum 10 (but any amount is okay) internships that are most relevant to the candidate's profile, based on what you know about them. 

        Avoid returning internships that are not relevant to the candidate's application range. For example, don't recommend master internships if you know explicitely they are not master students. 

        Return a JSON array containing only the selected internships, formatted exactly like this:
        [
            {{"company": "Company Name", "role": "Role Title", "location": "Location", "application_link"}}
        ] 

        Your response should only be JSON.
        """

        cleaned_response = ""  # Initialize outside try block
        try:
            response = recommendation_agent.run(prompt)
            print("Raw recommendation response:", response.content)
            # Clean the response
            cleaned_response = response.content.replace("```json", "").replace("```", "").strip()

            recommendations = json.loads(cleaned_response)

            # Process each recommendation to extract proper URLs and clean HTML
            processed_recommendations = []
            for job in recommendations:
                processed_job = {}

                # Extract company name from HTML if present
                if 'company' in job:
                    import re
                    company_match = re.search(r'<strong>(.*?)</strong>', job['company'])
                    if company_match:
                        processed_job['company'] = company_match.group(1)
                    else:
                        processed_job['company'] = re.sub(r'<[^>]+>', '', job['company'])

                # Copy role and location as is
                processed_job['role'] = job.get('role', '')
                processed_job['location'] = job.get('location', '')

                # Process application link/salary
                if 'application' in job or 'application_link' in job:
                    app_data = job.get('application') or job.get('application_link', '')
                    if app_data.startswith('$'):  # Handle salary
                        processed_job['application_link'] = app_data
                    else:  # Handle URL
                        url_match = re.search(r'href=["\'](.*?)["\']', app_data)
                        if url_match:
                            url = url_match.group(1)
                            # Clean up extracted URL
                            url = url.replace('https://https://', 'https://')
                            url = url.replace('http://https://', 'https://')
                            url = url.replace('https://http://', 'https://')
                            if not url.startswith(('http://', 'https://')):
                                url = 'https://' + url
                            processed_job['application_link'] = url
                        else:
                            processed_job['application_link'] = app_data.strip()

                processed_recommendations.append(processed_job)

            return processed_recommendations[:10]  # Ensure we only return max 10
        except Exception as e:
            print(f"Error in recommendations: {e}")
            print(f"Cleaned response that caused error: {cleaned_response}")
            return [internships[i] for i in range(min(10, len(internships)))]  # Fallback to first 10 if there's an error

# Configure Google AI
load_dotenv()

try:
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
    if not GOOGLE_API_KEY:
        print("⚠️ GOOGLE_API_KEY is missing in .env.")

    AGENT_QL_API_KEY = os.getenv("AGENT_QL_API_KEY")
    if AGENT_QL_API_KEY:
        agentql.configure(api_key=AGENT_QL_API_KEY)
    else:
        print("⚠️ AGENT_QL_API_KEY is missing in .env")

    genai.configure(api_key=GOOGLE_API_KEY)
except Exception as e:
    st.error("Missing API keys in .env")
    st.stop()
# Initialize the model
model = genai.GenerativeModel('gemini-2.0-pro-exp-02-05')


class JobSearchManager:
    """Handles job search operations"""

    def create_profile(self, resume_data: Dict, responses: Dict) -> Dict:
        """Creates a comprehensive profile from resume data and question responses"""
        return {
            "skills": resume_data.get("skills", []),
            "experiences": resume_data.get("experiences", []),
            "interests": resume_data.get("interests", []),
            "responses": responses,
            "metadata": {
                "creation_time": time.strftime("%Y-%m-%d %H:%M:%S"),
                "last_updated": time.strftime("%Y-%m-%d %H:%M:%S")
            }
        }

    def compile_job_links(self, resume_data: Dict, responses: Dict) -> Dict[str, List[Dict]]:
        """Process job search"""
        # Placeholder for new job search implementation
        return {
            "current_fit": [
                {
                    "title": "Example Position",
                    "company": "Example Corp",
                    "url": "https://example.com/jobs/1",
                    "location": "Remote, US"
                }
            ]
        }

class ResumeProcessor:
    def __init__(self):
        self.agent = Agent(
            model=Gemini(
                id="gemini-2.0-flash",
                api_key=GOOGLE_API_KEY
            ),
            description="You are an expert resume analyzer and career counselor.",
            markdown=True,
            tools=[GithubTool()]
        )
        self.database = {}
        self.github_tool = GithubTool()
        self.job_searcher = JobSearchManager()

    def parse_resume(self, file_content: str) -> Dict:
        # Existing method unchanged
        max_retries = 3
        for attempt in range(max_retries):
            prompt = f"Analyze this resume and extract key information including skills, experiences, and interests. This information will be provided to another artificial intelligence to generate follow up questions about the applicant, so you want to be very specific. Return the response as a JSON object with keys: skills, experiences, interests. In your JSON return do not include `````` at the end or start of the JSON. Only raw JSON, literally nothing else besides RAW JSON. You will create parsing issues, so only RAW JSON, no markdown code blocks whatsoever! Applicant Content: {file_content}"
            response = self.agent.run(prompt)
            print(f"Raw Gemini Response (Attempt {attempt + 1}):", response.content)

            if not response.content.strip():
                if attempt < max_retries - 1:
                    continue
                return {
                    "skills": [],
                    "experiences": [],
                    "interests": []
                }

            try:
                cleaned_content = response.content.replace("```json", "").replace("```", "").strip()
                return json.loads(cleaned_content)
            except json.JSONDecodeError as e:
                print(f"JSON Decode Error (Attempt {attempt + 1}):", e)
                if attempt < max_retries - 1:
                    continue
                return {
                    "skills": [],
                    "experiences": [],
                    "interests": []
                }

    def generate_questions(self, resume_data: Dict) -> List[str]:
        # Existing method unchanged
        prompt = f"""Based on this resume data: {json.dumps(resume_data)}
        Generate only 5 follow-up questions about:
        1. Role/Experience details about experiences that were in their resume. Specifically, use the STAR methodology to generate questions. You want to obtain information that you can later use to create cover letters for the applicant, so you want to be very specific. 
        2. Specific skills and their applications. You want to know specific times the applicant has applied their skills succesfully.
        3. Career interests and goals. This type of question is only to detail what type of ideal role the applicant projects themselves on or is working towards, both role-wise and skills-wise.  

        Consider that the applicant is looking for internships during the summer. These questions should ideally reflect some questions they may see interviewing or in application pages. Do not add Q1, Q2 at the beguinning of each question.  

        Return ONLY a JSON array of strings, formatted exactly like this: ["question1", "question2", "question3", "question4", "question5"]"""
        response = self.agent.run(prompt)
        print("Raw Questions Response:", response.content)
        try:
            cleaned_response = response.content.replace("```json", "").replace("```", "").strip()
            return json.loads(cleaned_response)
        except json.JSONDecodeError as e:
            print(f"JSON Decode Error: {e}")
            return [
                "Could you elaborate on your most recent role?",
                "What specific projects have you worked on using Python?",
                "How have you applied your data science skills in real projects?",
                "What are your career goals in the next 3-5 years?",
                "Which of your listed skills are you most eager to develop further?"
            ]

    def generate_search_terms(self, profile: Dict) -> Dict[str, List[str]]:
        # Existing method unchanged
        prompt = f"""Based on this applicant's profile
        {json.dumps(profile)}
        create a series of internships the applicant might be looking for (for example, data science intern, business analytics intern, etc) that match their profile, skill, demonstrated interests, and responses in the answers. Return ONLY a JSON array of strings, formatted exactly like this: ["Searchterm1", "Searchterm2", "Searchterm3", "Searchterm4", "Searchterm5"]"""
        response = self.agent.run(prompt)
        print("Raw Search Terms Response:", response.content)
        try:
            cleaned_content = response.content.replace("```json", "").replace("```", "").strip()
            return json.loads(cleaned_content)
        except json.JSONDecodeError as e:
            print(f"JSON Decode Error: {e}")
            return {
                "current_fit": ["Entry Level Position", "Junior Role"],
                "aspirational": ["Senior Position", "Team Lead"]
            }

    def enhanced_search(self, profile: Dict) -> Dict[str, List[Dict]]:
        """Public method to trigger job search"""
        resume_data = profile.get('resume_data', {})
        responses = profile.get('responses', {})
        return self.job_searcher.compile_job_links(resume_data, responses)

def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def read_docx(file):
    doc = docx.Document(file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def process_responses(processor):
    profile = {
        "resume_data": st.session_state['resume_data'],
        "responses": st.session_state.responses
    }

    # Generate job search terms
    search_terms = processor.generate_search_terms(profile)
    print("Recommended Job Search Terms:")


    # Store in database
    processor.database[st.session_state['resume_data'].get('name', 'Anonymous')] = profile
    return profile

def main():
    st.title("Internship Hunt w/Agno📚, Gemini 🧠, and AgentQL🦾")

    if 'initialized' not in st.session_state:
        st.session_state.update({
            'initialized': True,
            'form_submitted': False,
            'resume_processed': False,
            'responses': {},
            'questions': [],
            'resume_data': None,
            'processor': ResumeProcessor()
        })

    with st.form("resume_upload_form"):
        uploaded_file = st.file_uploader("Upload Resume", type=["pdf", "docx"])

        if st.form_submit_button("Analyze Resume") and uploaded_file:
            with st.spinner("Processing resume..."):
                file_content = read_pdf(uploaded_file) if uploaded_file.type == "application/pdf" else read_docx(uploaded_file)
                st.session_state.resume_data = st.session_state.processor.parse_resume(file_content)
                st.session_state.questions = st.session_state.processor.generate_questions(st.session_state.resume_data)
                st.session_state.resume_processed = True

    if st.session_state.resume_processed and not st.session_state.form_submitted:
        with st.form("responses_form"):
            st.subheader("Follow-up Questions")
            for i, question in enumerate(st.session_state.questions):
                key = f"q_{i}"
                if key not in st.session_state.responses:
                    st.session_state.responses[key] = ""
                st.session_state.responses[key] = st.text_area(
                    label=f"Q{i+1}: {question}",
                    value=st.session_state.responses[key],
                    key=key
                )

            if st.form_submit_button("Submit Responses"):
                with st.spinner("Processing your answers..."):
                    st.session_state.form_submitted = True
                    profile = process_responses(st.session_state.processor)
                    st.session_state.job_links = st.session_state.processor.enhanced_search(profile)
                st.rerun()

    if st.session_state.form_submitted:
        st.success("Responses saved successfully!")
        for i, question in enumerate(st.session_state.questions):
            st.text_area(
                label=f"Q{i+1}: {question}",
                value=st.session_state.responses.get(f"q_{i}", ""),
                disabled=True,
                key=f"readonly_{i}"
            )

        st.subheader("Recommended Internships")

        with st.spinner("Fetching and analyzing internship opportunities..."):
            # Fetch internships from GitHub
            github_tool = GithubTool()
            internship_data = asyncio.run(github_tool.fetch_internship_data())

        if internship_data["status"] == "success":
            profile = {
                "resume_data": st.session_state.resume_data,
                "responses": st.session_state.responses
            }

            with st.spinner("Analyzing internships for best matches..."):
                recommended_internships = github_tool.recommend_internships(
                    internship_data["internships"], 
                    profile
                )

            # Cleanup session
            asyncio.run(github_tool.close_session())

            # Filter and clean internships with valid application links
            valid_internships = []
            for job in recommended_internships:
                if 'application_link' in job and job['application_link']:
                    application_url = job['application_link'].strip()
                    # Skip if the application link is just a salary
                    if application_url.startswith('$'):
                        continue

                    # Clean application URL if needed
                    if 'href=' in application_url:
                        import re
                        url_match = re.search(r'href=["\'](.*?)["\']', application_url)
                        if url_match:
                            application_url = url_match.group(1).strip()

                    # Remove any accidentally prepended protocols
                    application_url = application_url.replace('https://https://', 'https://')
                    application_url = application_url.replace('http://https://', 'https://')
                    application_url = application_url.replace('https://http://', 'https://')

                    import re
                    # Clean company name from HTML tags
                    if 'href=' in job.get('company', ''):
                        company_match = re.search(r'<strong>(.*?)</strong>', job['company'])
                        if company_match:
                            job['company'] = company_match.group(1)
                        else:
                            # Fallback: strip all HTML tags
                            job['company'] = re.sub(r'<[^>]+>', '', job['company'])

                    # Ensure URL has proper protocol
                    if not application_url.startswith(('http://', 'https://')):
                        application_url = 'https://' + application_url

                    job['application_link'] = application_url
                    valid_internships.append(job)

            if valid_internships:
                # Create placeholder for internships
                internship_containers = {}

                for job in valid_internships:
                    application_url = job.get('application_link', '')
                    if (application_url and 
                        application_url.startswith(('http://', 'https://')) and 
                        not application_url.startswith('$')):

                        # Create a container for this job
                        container = st.container()
                        internship_containers[application_url] = container

                        with container:
                            st.markdown(f"### {job['company']} - {job['role']}")
                            st.markdown(f"📍 {job['location']}")
                            with st.spinner("Analyzing application requirements..."):

                                # Analyze requirements for this specific internship
                                requirements, should_display = analyze_application_link(application_url)
                                if not should_display:
                                    continue

                                job['requirements'] = requirements

                                # Create three columns for requirement indicators
                                col1, col2, col3 = st.columns(3)

                                # Display cover letter requirement status
                                requires_cover_letter = requirements.get('requires_cover_letter', False)
                                if requires_cover_letter:
                                    st.info("📝 Cover Letter Required")
                                    cover_letter_key = f"cover_letter_{job['company']}_{job['role']}"
                                    if cover_letter_key not in st.session_state:
                                        st.session_state[cover_letter_key] = ""
                                else:
                                    st.success("✅ No Cover Letter Required")

                                # Display research statement requirement status
                                requires_research_statement = requirements.get('requires_research_statement', False)
                                if requires_research_statement:
                                    st.info("🔬 Research Statement Required")
                                    research_statement_key = f"research_statement_{job['company']}_{job['role']}"
                                    if research_statement_key not in st.session_state:
                                        st.session_state[research_statement_key] = ""
                                else:
                                    st.success("✅ No Research Statement Required")

                                # Display "why us" statement requirement status
                                requires_why_us_statement = requirements.get('requires_why_us_statement', False)
                                if requires_why_us_statement:
                                    st.info("🎯 'Why Us' Statement Required")
                                    why_us_statement_key = f"why_us_statement_{job['company']}_{job['role']}"
                                    if why_us_statement_key not in st.session_state:
                                        st.session_state[why_us_statement_key] = ""
                                else:
                                    st.success("✅ No 'Why Us' Statement Required")

                            # Display apply button
                            st.markdown(f"<a href='{job['application_link']}' target='_blank'>🔗 Apply Now</a>", unsafe_allow_html=True)

                            # Generate cover letter if requirements say it's needed
                            if requirements.get('requires_cover_letter'):
                                # Create unique session state key for this job
                                cover_letter_key = f"cover_letter_{job['company']}_{job['role']}"

                                # Auto-generate cover letter if not already generated
                                if cover_letter_key not in st.session_state or not st.session_state[cover_letter_key]:
                                    with st.spinner("Generating your cover letter..."):
                                        st.session_state[cover_letter_key] = generate_cover_letter(
                                            {"resume_data": st.session_state.resume_data, 
                                             "responses": st.session_state.responses}, 
                                            job
                                        )

                                # Display cover letter
                                st.text_area(
                                    "Generated Cover Letter",
                                    value=st.session_state[cover_letter_key],
                                    height=200,
                                    key=f"cl_display_{job['company']}_{job['role']}"
                                )

                                # Generate download links
                                dl_col1_cl, dl_col2_cl = st.columns(2)

                                # Store files in session state
                                files_key_cl = f"files_cl_{job['company']}_{job['role']}"
                                if files_key_cl not in st.session_state:
                                    profile_data = {
                                        "resume_data": st.session_state.resume_data
                                    }
                                    st.session_state[files_key_cl] = {
                                        'docx': create_docx_cover_letter(
                                            st.session_state[cover_letter_key],
                                            job['company'],
                                            job['role'],
                                            profile_data
                                        )
                                    }

                                with dl_col1_cl:
                                    company_safe = "".join(c for c in job['company'] if c.isalnum())
                                    role_safe = "".join(c for c in job['role'] if c.isalnum())
                                    st.markdown(
                                        f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(st.session_state[files_key_cl]["docx"]).decode()}" download="cover_letter_{company_safe}_{role_safe}.docx" target="_blank">📝 Cover Letter (DOCX)</a>',
                                        unsafe_allow_html=True
                                    )

                            # Generate research statement if required
                            if requirements.get('requires_research_statement'):
                                # Create unique session state key
                                research_statement_key = f"research_statement_{job['company']}_{job['role']}"

                                # Auto-generate research statement if not already generated
                                if research_statement_key not in st.session_state or not st.session_state[research_statement_key]:
                                    with st.spinner("Generating your research statement..."):
                                        st.session_state[research_statement_key] = generate_research_statement(
                                            {"resume_data": st.session_state.resume_data,
                                             "responses": st.session_state.responses},
                                            job
                                        )

                                # Display research statement
                                st.text_area(
                                    "Generated Research Statement",
                                    value=st.session_state[research_statement_key],
                                    height=200,
                                    key=f"rs_display_{job['company']}_{job['role']}"
                                )

                                # Generate download links for research statement
                                dl_col1_rs, dl_col2_rs = st.columns(2)
                                files_key_rs = f"files_rs_{job['company']}_{job['role']}"
                                if files_key_rs not in st.session_state:
                                    profile_data = {
                                        "resume_data": st.session_state.resume_data
                                    }
                                    st.session_state[files_key_rs] = {
                                        'docx': create_docx_cover_letter(
                                            st.session_state[research_statement_key],
                                            job['company'],
                                            job['role'],
                                            profile_data
                                        )
                                    }

                                with dl_col1_rs:
                                    company_safe = "".join(c for c in job['company'] if c.isalnum())
                                    role_safe = "".join(c for c in job['role'] if c.isalnum())
                                    st.markdown(
                                        f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(st.session_state[files_key_rs]["docx"]).decode()}" download="research_statement_{company_safe}_{role_safe}.docx" target="_blank">📝 Research Statement (DOCX)</a>',
                                        unsafe_allow_html=True
                                    )                                

                            # Generate "why us" statement if required
                            if requirements.get('requires_why_us_statement'):
                                # Create unique session state key
                                why_us_statement_key = f"why_us_statement_{job['company']}_{job['role']}"

                                # Auto-generate "why us" statement if not already generated
                                if why_us_statement_key not in st.session_state or not st.session_state[why_us_statement_key]:
                                    with st.spinner("Generating your 'Why Us' statement..."):
                                        st.session_state[why_us_statement_key] = generate_why_us_statement(
                                            {"resume_data": st.session_state.resume_data,
                                             "responses": st.session_state.responses},
                                            job
                                        )

                                # Display "why us" statement
                                st.text_area(
                                    "Generated 'Why Us' Statement",
                                    value=st.session_state[why_us_statement_key],
                                    height=200,
                                    key=f"wu_display_{job['company']}_{job['role']}"
                                )

                                # Generate download links for "why us" statement
                                dl_col1_wu, dl_col2_wu = st.columns(2)
                                files_key_wu = f"files_wu_{job['company']}_{job['role']}"
                                if files_key_wu not in st.session_state:
                                    profile_data = {
                                        "resume_data": st.session_state.resume_data
                                    }
                                    st.session_state[files_key_wu] = {
                                        'docx': create_docx_cover_letter(
                                            st.session_state[why_us_statement_key],
                                            job['company'],
                                            job['role'],
                                            profile_data
                                        )
                                    }

                                with dl_col1_wu:
                                    company_safe = "".join(c for c in job['company'] if c.isalnum())
                                    role_safe = "".join(c for c in job['role'] if c.isalnum())
                                    st.markdown(
                                        f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(st.session_state[files_key_wu]["docx"]).decode()}" download="why_us_statement_{company_safe}_{role_safe}.docx" target="_blank">📝 Why Us Statement (DOCX)</a>',
                                        unsafe_allow_html=True
                                    )
                            
                            st.divider()

if __name__ == "__main__":
    main()